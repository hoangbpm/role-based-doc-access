from docx import Document
from docx.shared import Inches, Pt
from datetime import datetime

# Mở file Word
file_path = r"C:\Users\ASUS\Documents\BTL_thiết kế hướng đối tượng\Report _group26_2.docx"  # Đổi thành file của bạn
doc = Document(file_path)

# Đường dẫn ảnh chữ ký
signature_path = r"C:\Users\ASUS\Documents\BTL_thiết kế hướng đối tượng\Screenshot 2025-02-28 192944.png"  # Thay bằng đường dẫn ảnh chữ ký

# Thông tin giám đốc
director_name = "Nguyễn Văn V"
current_date = datetime.now().strftime("Ngày %d/%m/%Y")


# Tìm dòng "Chữ ký giám đốc"
for i, para in enumerate(doc.paragraphs):
    if "Chữ ký giám đốc" in para.text:
        print(f"✅ Tìm thấy 'Chữ ký giám đốc' ở đoạn {i}")

        # Tạo đoạn mới để căn giữa chữ ký + ngày tháng
        new_para = doc.add_paragraph()
        new_para.alignment = 1  # Căn giữa (0: trái, 1: giữa, 2: phải)

        # Thêm ngày tháng trước
        run_date = new_para.add_run(current_date)
        run_date.font.size = Pt(12)  # Cỡ chữ 12

        # Thêm ảnh chữ ký
        new_para.add_run("\n")  # Xuống dòng trước khi chèn ảnh
        run_sig = new_para.add_run()
        run_sig.add_picture(signature_path, width=Inches(1.5))

        # Thêm tên giám đốc
        new_para.add_run("\n")  # Xuống dòng
        run_name = new_para.add_run(director_name)
        run_name.font.size = Pt(12)
        run_name.bold = True  # In đậm

        break  # Dừng sau khi chèn xong

# Lưu file mới
doc.save("file_da_co_chu_ky.docx")

print("✅ Đã chèn chữ ký, căn giữa và định dạng đúng!")
